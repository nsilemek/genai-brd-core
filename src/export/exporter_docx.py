from __future__ import annotations

import os
from typing import Any, Dict, Optional

from docx import Document
from docx.shared import Pt


def _add_kv_section(doc: Document, title: str, value: Any) -> None:
    doc.add_heading(title, level=2)

    if isinstance(value, list):
        if value:
            for item in value:
                doc.add_paragraph(str(item), style="List Bullet")
        else:
            doc.add_paragraph("(empty)")
        return

    text = str(value or "").strip()
    doc.add_paragraph(text if text else "(empty)")


def export_docx_file(
    out_dir: str,
    session_id: str,
    fields: Dict[str, Any],
    scores: Optional[Dict[str, Any]] = None,
    filename: Optional[str] = None,
    title: str = "BRD / TO-BE JOURNEY",
) -> str:
    os.makedirs(out_dir, exist_ok=True)
    if not filename:
        filename = f"brd_{session_id}.docx"
    path = os.path.join(out_dir, filename)

    doc = Document()
    doc.add_heading(title, level=1)

    # Optional meta
    doc.add_paragraph(f"Session ID: {session_id}")

    doc.add_paragraph("")  # spacer

    order = [
        "Background",
        "Expected Results",
        "Target Customer Group",
        "Impacted Channels",
        "Impacted Journey",
        "Journeys Description",
        "Reports Needed",
        "Traffic Forecast",
    ]

    for k in order:
        _add_kv_section(doc, k, fields.get(k, ""))

    # Scores section
    if scores:
        doc.add_page_break()
        doc.add_heading("Score Summary", level=1)
        doc.add_paragraph(f"Total Score: {scores.get('total_score')} / {scores.get('max_total')}")
        doc.add_paragraph(f"Submit Allowed: {scores.get('submit_allowed')}")

        blockers = scores.get("submit_blockers") or []
        if blockers:
            doc.add_heading("Blockers", level=2)
            for b in blockers:
                doc.add_paragraph(str(b), style="List Bullet")

        weak_fields = scores.get("weak_fields") or []
        if weak_fields:
            doc.add_heading("Weak Fields", level=2)
            for wf in weak_fields:
                doc.add_paragraph(str(wf), style="List Bullet")

    # Light typography tweak (optional)
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    doc.save(path)
    return path
